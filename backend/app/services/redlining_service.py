"""Contract redlining and risk analysis service"""
import structlog
from typing import List, Dict
from docx import Document
from docx.shared import RGBColor

logger = structlog.get_logger()

class RedliningService:
    """Service for analyzing contracts and marking risky clauses"""

    def __init__(self):
        self.risk_keywords = {
            "high": [
                "безоговорочно", "без права", "в одностороннем порядке",
                "не подлежит", "отказывается от", "освобождается от ответственности"
            ],
            "medium": [
                "по своему усмотрению", "вправе изменить", "имеет право отказать"
            ],
            "low": ["может", "рекомендуется", "желательно"]
        }

    async def analyze_contract(self, docx_path: str, party_role: str = "заказчик", risk_level: str = "medium") -> Dict:
        """Analyze contract for risks"""
        try:
            doc = Document(docx_path)
            risks = []
            risk_score = 0

            for i, para in enumerate(doc.paragraphs):
                text = para.text.lower()
                if not text.strip():
                    continue

                found_risks = self._find_risks_in_text(text, risk_level)
                for risk in found_risks:
                    risks.append({
                        "id": f"risk_{i}_{len(risks)}",
                        "text_range": f"Параграф {i+1}",
                        "original_text": para.text[:100],
                        "suggestion": risk["suggestion"],
                        "risk_type": risk["type"],
                        "legal_basis": "ГК РК ст. 380-385"
                    })
                    risk_score += risk["score"]

            risk_score = min(100, risk_score)
            return {
                "risks": risks,
                "risk_score": risk_score,
                "summary": self._generate_summary(len(risks), risk_score, party_role),
                "party_role": party_role
            }
        except Exception as e:
            logger.error("contract_analysis_failed", error=str(e))
            raise

    def _find_risks_in_text(self, text: str, sensitivity: str) -> List[Dict]:
        risks = []
        levels = ["high"]
        if sensitivity in ["medium", "high"]:
            levels.append("medium")
        if sensitivity == "high":
            levels.append("low")

        for level in levels:
            for keyword in self.risk_keywords[level]:
                if keyword in text:
                    risks.append({
                        "type": level,
                        "suggestion": f"Обратите внимание на '{keyword}'",
                        "score": {"high": 15, "medium": 8, "low": 3}[level]
                    })
        return risks

    def _generate_summary(self, risks_count: int, risk_score: int, party_role: str) -> str:
        if risk_score >= 70:
            verdict = "ВЫСОКИЙ РИСК"
        elif risk_score >= 40:
            verdict = "СРЕДНИЙ РИСК"
        else:
            verdict = "НИЗКИЙ РИСК"
        return f"Для {party_role}: найдено {risks_count} рисков, оценка {risk_score}/100 - {verdict}"

    def create_redlined_document(self, input_path: str, output_path: str, risks: List[Dict], party_role: str):
        try:
            doc = Document(input_path)
            for risk in risks[:20]:
                para_idx = self._extract_para_index(risk["text_range"])
                if para_idx and para_idx < len(doc.paragraphs):
                    para = doc.paragraphs[para_idx]
                    para.add_run(f"\n[{risk['risk_type'].upper()}] {risk['suggestion']}").font.color.rgb = RGBColor(220, 20, 60)
            doc.save(output_path)
        except Exception as e:
            logger.error("redlining_failed", error=str(e))
            raise

    def _extract_para_index(self, text_range: str):
        try:
            parts = text_range.split()
            if len(parts) >= 2:
                return int(parts[1]) - 1
        except:
            pass
        return None

_redlining_service = None

def get_redlining_service():
    global _redlining_service
    if _redlining_service is None:
        _redlining_service = RedliningService()
    return _redlining_service
